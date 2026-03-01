from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# スタイル設定
style = doc.styles['Normal']
style.font.size = Pt(11)

# タイトル
title = doc.add_heading('レストラン来客数予測', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Kaggle Recruit Restaurant Visitor Forecasting')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('発表時間目安: 10-15分')
doc.add_paragraph()

# 1. イントロダクション
doc.add_heading('1. イントロダクション', level=1)
doc.add_heading('スライド1: タイトル', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('本日は、Kaggleコンペティション「Recruit Restaurant Visitor Forecasting」を題材に、機械学習モデルの比較検証を行った結果を発表します。')
doc.add_paragraph('このコンペでは、日本のレストランの来客数を予測します。データは、ホットペッパーグルメとAirレジという2つのシステムから提供されています。')

# 2. 問題設定
doc.add_heading('2. 問題設定', level=1)
doc.add_heading('スライド2: なぜ来客予測が重要か', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('レストランにとって、来客数の予測は非常に重要です。')
doc.add_paragraph('・食材の仕入れ: 多すぎれば廃棄、少なすぎれば品切れ')
doc.add_paragraph('・スタッフ配置: 適切な人員配置でコスト削減')
doc.add_paragraph('・顧客満足度: 待ち時間の短縮')

doc.add_heading('スライド3: 評価指標 RMSLE', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('評価指標はRMSLE、Root Mean Squared Logarithmic Errorです。')
doc.add_paragraph('この指標の特徴は、過少予測をより厳しく罰することです。')
doc.add_paragraph('レストランにとって、食材が足りなくなる（過少予測）方が、余る（過大予測）より深刻だからです。')

# 3. データの事前処理
doc.add_heading('3. データの事前処理', level=1)
doc.add_heading('スライド4: データ概要', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('データは8つのファイルから構成されています。')
doc.add_paragraph('・air_visit_data.csv: 来客履歴（約25万件）')
doc.add_paragraph('・air_store_info.csv: 店舗情報（ジャンル、エリア）')
doc.add_paragraph('・date_info.csv: 祝日情報')
doc.add_paragraph('期間は2016年1月から2017年4月までの約16ヶ月分です。')

doc.add_heading('スライド5: データの問題点と対処', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

p = doc.add_paragraph()
p.add_run('問題1: 欠損日').bold = True
doc.add_paragraph('営業していない日はデータがありません。')
doc.add_paragraph('→ 対処: 日付でリサンプルし、欠損日は0で補完')

p = doc.add_paragraph()
p.add_run('問題2: 外れ値').bold = True
doc.add_paragraph('一部の店舗で異常に多い来客数')
doc.add_paragraph('→ 対処: 2.4σを超える値をクリップ')

p = doc.add_paragraph()
p.add_run('問題3: 新規店舗').bold = True
doc.add_paragraph('履歴が少ない店舗は予測が困難')
doc.add_paragraph('→ 対処: 同じジャンル・エリアの平均値で補完')

doc.add_heading('スライド6: 特徴量エンジニアリング', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('上位解法を分析した結果、特徴量エンジニアリングが最も重要だとわかりました。')
doc.add_paragraph('1. 時間特徴量: 曜日、月、祝日フラグ')
doc.add_paragraph('2. Rolling統計量: 過去7日、14日、21日の平均・標準偏差')
doc.add_paragraph('3. ラグ特徴量: 1週間前、2週間前の来客数')
doc.add_paragraph('4. 指数加重平均: 直近のデータを重視した平均')

# 4. ランダムフォレスト
doc.add_heading('4. ランダムフォレストによる分析', level=1)
doc.add_heading('スライド7: ランダムフォレストとは', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('まず、ベースラインモデルとしてランダムフォレストを使用しました。')
doc.add_paragraph('ランダムフォレストは、複数の決定木を作成し、その平均を取る手法です。')
doc.add_paragraph('メリット: 外れ値に強い、スケーリング不要、過学習しにくい')

doc.add_heading('スライド8: ランダムフォレストの結果', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('・訓練データ RMSLE: 約0.45')
doc.add_paragraph('・検証データ RMSLE: 約0.54')
doc.add_paragraph('一定の精度は出ましたが、上位解法と比較すると劣っています。')

doc.add_heading('スライド9: ランダムフォレストの問題点', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

p = doc.add_paragraph()
p.add_run('問題1: 計算速度が遅い').bold = True
doc.add_paragraph('100本の決定木を学習するのに数分かかります。')

p = doc.add_paragraph()
p.add_run('問題2: メモリ消費が大きい').bold = True
doc.add_paragraph('全ての決定木をメモリに保持する必要があります。')

p = doc.add_paragraph()
p.add_run('問題3: 精度の限界').bold = True
doc.add_paragraph('勾配ブースティング系と比較して精度が劣ります。')
doc.add_paragraph('結論: ベースラインとしては有用だが、最終モデルには不向き')

# 5. XGBoost
doc.add_heading('5. XGBoostによる分析', level=1)
doc.add_heading('スライド10: XGBoostとは', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('XGBoostは勾配ブースティングの代表的な実装で、2014年頃からKaggleで広く使われています。')
doc.add_paragraph('特徴: 正則化による過学習防止、欠損値の自動処理、並列処理対応')

doc.add_heading('スライド11: XGBoostの結果', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('・訓練データ RMSLE: 約0.42')
doc.add_paragraph('・検証データ RMSLE: 約0.52')
doc.add_paragraph('ランダムフォレストより明確に改善しました。')

# 6. LightGBM
doc.add_heading('6. LightGBMによる分析', level=1)
doc.add_heading('スライド12: LightGBMとは', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('LightGBMはMicrosoftが開発した勾配ブースティングで、上位解法の大半が採用しています。')
doc.add_paragraph('なぜLightGBMか:')
doc.add_paragraph('・XGBoostの10倍以上高速')
doc.add_paragraph('・メモリ効率が良い')
doc.add_paragraph('・カテゴリ変数を直接扱える')

doc.add_heading('スライド13: LightGBMの結果', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('・訓練データ RMSLE: 約0.40')
doc.add_paragraph('・検証データ RMSLE: 約0.50')
doc.add_paragraph('3つのモデル中、最も良い結果となりました。')

# 7. モデル比較
doc.add_heading('7. モデル比較まとめ', level=1)
doc.add_heading('スライド14: 比較表', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'モデル'
hdr[1].text = '検証RMSLE'
hdr[2].text = '学習時間'
hdr[3].text = 'メモリ'
row1 = table.rows[1].cells
row1[0].text = 'Random Forest'
row1[1].text = '0.54'
row1[2].text = '遅い'
row1[3].text = '大'
row2 = table.rows[2].cells
row2[0].text = 'XGBoost'
row2[1].text = '0.52'
row2[2].text = '中程度'
row2[3].text = '中'
row3 = table.rows[3].cells
row3[0].text = 'LightGBM'
row3[1].text = '0.50'
row3[2].text = '速い'
row3[3].text = '小'

doc.add_paragraph()
doc.add_paragraph('LightGBMが精度・速度・メモリ全てで優れています。')

# 8. 結論
doc.add_heading('8. 結論と学び', level=1)
doc.add_heading('スライド16: 結論', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('1. 特徴量エンジニアリングが最重要')
doc.add_paragraph('   - モデル選択より、良い特徴量を作ることが重要')
doc.add_paragraph('   - Rolling統計量が最も効果的')
doc.add_paragraph('2. LightGBMが最適解')
doc.add_paragraph('   - 精度、速度、メモリ効率の全てで優秀')
doc.add_paragraph('3. 時系列データの扱いに注意')
doc.add_paragraph('   - ランダム分割は禁止、必ず時間ベースで分割する')

# 9. 改善点
doc.add_heading('9. 改善点とその理由 - なぜ精度が上がったか', level=1)

doc.add_heading('スライド17: 精度改善の全体像', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True
doc.add_paragraph('改善は大きく2つの観点から説明できます：')
doc.add_paragraph('1. アルゴリズムの違い（バギング vs ブースティング）')
doc.add_paragraph('2. 実装の最適化（XGBoost vs LightGBM）')

doc.add_heading('スライド18: バギング vs ブースティング', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

p = doc.add_paragraph()
p.add_run('ランダムフォレスト（バギング）:').bold = True
doc.add_paragraph('各木は独立して学習 → 予測を平均化')

p = doc.add_paragraph()
p.add_run('勾配ブースティング（XGBoost/LightGBM）:').bold = True
doc.add_paragraph('木1 → 誤差計算 → 木2が誤差を修正 → 木3がさらに修正...')

doc.add_paragraph('なぜブースティングが優れるか:')
doc.add_paragraph('・誤差を段階的に修正するため、より細かいパターンを学習可能')
doc.add_paragraph('・難しいサンプルに集中的に対応できる')

doc.add_heading('スライド20: XGBoostの改善ポイント', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

p = doc.add_paragraph()
p.add_run('改善点1: 勾配ベースの学習').bold = True
doc.add_paragraph('前の木の予測誤差（勾配）を次の木が学習 → 誤差を効率的に減らせる')

p = doc.add_paragraph()
p.add_run('改善点2: 正則化').bold = True
doc.add_paragraph('L1/L2正則化により過学習を防止 → 汎化性能が向上')

p = doc.add_paragraph()
p.add_run('改善点3: 欠損値の自動処理').bold = True
doc.add_paragraph('最適な分岐方向を自動で決定 → 情報を失わない')

doc.add_paragraph('結果: RMSLE 0.54 → 0.52（約4%改善）')

doc.add_heading('スライド21: LightGBMの改善ポイント', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

p = doc.add_paragraph()
p.add_run('改善点1: Leaf-wise成長戦略').bold = True
doc.add_paragraph('XGBoost: Level-wise（層ごとに均等に成長）')
doc.add_paragraph('LightGBM: Leaf-wise（損失が大きい葉を優先成長）')
doc.add_paragraph('→ 同じ計算量でより深い学習が可能')

p = doc.add_paragraph()
p.add_run('改善点2: ヒストグラムベースのアルゴリズム').bold = True
doc.add_paragraph('連続値を離散化（ビン化）して計算を高速化')
doc.add_paragraph('→ 10倍以上高速化、精度はほぼ同等')

p = doc.add_paragraph()
p.add_run('改善点3: カテゴリ変数の直接処理').bold = True
doc.add_paragraph('One-hotエンコーディング不要')
doc.add_paragraph('→ メモリ効率が向上')

doc.add_paragraph('結果: RMSLE 0.52 → 0.50（約4%改善）')

doc.add_heading('スライド22: 精度改善のまとめ', level=2)
p = doc.add_paragraph()
p.add_run('原稿:').bold = True

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '改善ステップ'
hdr2[1].text = 'RMSLE'
hdr2[2].text = '改善幅'
hdr2[3].text = '主な理由'
r1 = table2.rows[1].cells
r1[0].text = 'Random Forest'
r1[1].text = '0.54'
r1[2].text = '-'
r1[3].text = 'ベースライン'
r2 = table2.rows[2].cells
r2[0].text = '→ XGBoost'
r2[1].text = '0.52'
r2[2].text = '-0.02'
r2[3].text = '勾配ブースティング'
r3 = table2.rows[3].cells
r3[0].text = '→ LightGBM'
r3[1].text = '0.50'
r3[2].text = '-0.02'
r3[3].text = 'Leaf-wise + ヒストグラム'

doc.add_paragraph()
doc.add_paragraph('総改善: 0.54 → 0.50（約7.4%改善）')

# Q&A
doc.add_heading('補足: Q&A想定', level=1)

p = doc.add_paragraph()
p.add_run('Q: なぜランダムフォレストは精度が劣るのか？').bold = True
doc.add_paragraph('A: ランダムフォレストは各木が独立して学習するバギング手法です。一方、勾配ブースティングは前の木の誤差を修正しながら学習するため、より精度が高くなります。')

p = doc.add_paragraph()
p.add_run('Q: LightGBMが速い理由は？').bold = True
doc.add_paragraph('A: Histogram-basedアルゴリズムとLeaf-wise成長戦略を採用しているためです。データを離散化することで計算量を削減しています。')

p = doc.add_paragraph()
p.add_run('Q: 実務でもLightGBMを使うべきか？').bold = True
doc.add_paragraph('A: テーブルデータの回帰・分類タスクでは、LightGBMは非常に有力な選択肢です。ただし、解釈性が重要な場合は線形モデルや決定木も検討すべきです。')

# 保存
output_path = 'competitions/recruit-restaurant-visitor-forecasting/docs/presentation_script.docx'
doc.save(output_path)
print(f'Word文書を作成しました: {output_path}')
